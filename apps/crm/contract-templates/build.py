# -*- coding: utf-8 -*-
"""
Gerador dos templates .docx de contrato da Studio Sal (mentoria, assessoria,
branding_pessoal). Produz os .docx canônicos neste mesmo diretório, que são:
  - versionados no repo (fonte de verdade + fixtures dos testes de render);
  - subidos pro Supabase Storage por scripts/upload-contract-templates.ts.

O texto ESTÁTICO do contrato (cláusulas, bloco de assinatura, dados fixos da
CONTRATADA) vive aqui, no template. O conteúdo DINÂMICO (prazo, pagamento,
qualificação PF/PJ do contratante, valores) é resolvido no código em
src/server/lib/contract-data-builder e entra como {placeholder} único —
docxtemplater faz o merge no download.

Regenerar (raro — normalmente o dono edita o .docx no Word e sobe em /admin):
  python3 -m venv .venv && .venv/bin/pip install python-docx
  .venv/bin/python build.py
Depois subir pro Storage: pnpm --filter crm exec tsx scripts/upload-contract-templates.ts

Placeholders disponíveis (preenchidos pelo contract-data-builder):
  {nome} {nome_completo} {cpf_cnpj} {rg} {email} {whatsapp} {produto}
  {valor} {valor_extenso} {prazo} {pagamento} {qualificacao_contratante}
  {endereco} {endereco_logradouro} {endereco_numero} {endereco_complemento}
  {endereco_bairro} {endereco_cidade} {endereco_estado} {endereco_cep} {data}
Campo sem dado sai vazio (nullGetter no render) — nunca quebra a geração.
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# Dados fixos da CONTRATADA (config da própria Studio Sal — não é dado dinâmico,
# vive aqui no template, não no core).
CONTRATADA_BLOCK = (
    "Studio Sal LTDA, agência boutique especializada em branding pessoal, "
    "inscrita no CNPJ sob o nº 55.217.203/0001-26, com sede na Rua Belgrado, "
    "nº 65, andar 1, conjunto 11, bairro Vila Moinho Velho, São Paulo/SP, "
    "CEP 04.285-040, neste ato representada por Giulia Salvatore Tebet "
    "Moreira, brasileira, empresária, portadora do RG nº 39.287.985-2 e do "
    "CPF nº 041.422.198-01, sendo os contatos oficiais realizados pelo "
    "e-mail giulia@studiosal.com.br, doravante denominada STUDIO"
)


def heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(14)
    doc.add_paragraph()


def clause(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)


def body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(10)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)


def resumo_table(doc, objeto):
    t = doc.add_table(rows=5, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Tema", "Descrição"),
        ("Partes", "{nome_completo} e Studio Sal LTDA"),
        ("Objeto", objeto),
        ("Valor", "{valor} ({valor_extenso}) — {pagamento}"),
        ("Prazo", "{prazo}"),
    ]
    for i, (a, b) in enumerate(rows):
        c = t.rows[i].cells
        c[0].paragraphs[0].add_run(a).bold = i == 0
        c[1].paragraphs[0].add_run(b).bold = i == 0
    doc.add_paragraph()


def build(titulo, objeto_resumo, objeto_texto, servicos_fn, out_path, nao_participacao=False):
    doc = Document()
    s = doc.sections[0]
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

    heading(doc, titulo)
    body(doc, "Data: {data} | Local: São Paulo")

    clause(doc, "Resumo Executivo")
    resumo_table(doc, objeto_resumo)

    clause(doc, "PARTES")
    body(
        doc,
        "CONTRATANTE: {qualificacao_contratante}, doravante denominada CLIENTE; "
        f"e CONTRATADA: {CONTRATADA_BLOCK}. Têm entre si justo e acertado o "
        "presente contrato, que se regerá pelas cláusulas a seguir.",
    )

    clause(doc, "1. DO OBJETO")
    body(doc, objeto_texto)
    body(
        doc,
        "Parágrafo único: sendo este contrato uma obrigação de meio, a STUDIO "
        "não se responsabiliza pelos resultados obtidos pela CLIENTE.",
    )

    clause(doc, "2. DOS SERVIÇOS")
    servicos_fn(doc)

    clause(doc, "3. RESPONSABILIDADES DAS PARTES")
    body(doc, "A STUDIO compromete-se a: executar os serviços com zelo e qualidade; manter sigilo sobre as informações da CLIENTE; cumprir o objeto e os serviços descritos nas Cláusulas 1ª e 2ª.")
    body(doc, "A CLIENTE compromete-se a: efetuar os pagamentos nos prazos acordados; fornecer dados verdadeiros e os materiais/acessos necessários; participar ativamente das atividades propostas; fazer uso pessoal e não comercial do conteúdo/metodologia.")

    clause(doc, "4. VALOR E PAGAMENTO")
    body(doc, "4.1. Pelo objeto deste contrato, a CLIENTE pagará à STUDIO o valor total de {valor} ({valor_extenso}), {pagamento}.")
    body(doc, "4.2. A não quitação nos prazos pactuados implica suspensão dos serviços, multa de 10% (dez por cento) e juros de 1% (um por cento) ao mês.")
    body(doc, "4.3. Os pagamentos devem ser comprovados por e-mail ou WhatsApp em até 24 horas — é responsabilidade da CLIENTE essa comprovação.")
    body(doc, "4.4. Em caso de atraso superior a 30 (trinta) dias, o contrato poderá ser rescindido por justa causa pela STUDIO, retidos os materiais já elaborados até a quitação.")
    if nao_participacao:
        # Verbatim do contrato-modelo (Mentoria Salto) — refere encontros/plataforma.
        body(doc, "4.5. A não participação da CLIENTE nas atividades propostas, nos encontros ou na utilização da plataforma, não será motivo de redução no valor do contrato, independentemente de sua motivação.")

    clause(doc, "5. VIGÊNCIA")
    body(doc, "5.1. Este contrato é válido por {prazo} contados da data de assinatura. O encerramento da vigência não exclui a obrigação de pagar parcelas vencidas posteriormente.")
    body(doc, "5.2. Nos termos do art. 49 do Código de Defesa do Consumidor, a CLIENTE poderá desistir no prazo de 7 (sete) dias corridos a contar da assinatura ou do primeiro acesso, com devolução integral. Após esse prazo, não haverá reembolso dos valores já pagos.")

    clause(doc, "6. RESCISÃO")
    body(doc, "6.1. O contrato pode ser encerrado por qualquer das partes mediante aviso prévio de 30 (trinta) dias.")
    body(doc, "6.2. No encerramento antecipado pela CLIENTE, esta arcará com o valor proporcional aos dias já decorridos, além de multa compensatória de 50% (cinquenta por cento) do valor restante contratado.")

    clause(doc, "7. DIREITO DE USO E DIVULGAÇÃO")
    body(doc, "A CLIENTE autoriza, a título gratuito, o uso de seu nome, imagem e depoimentos pela STUDIO para divulgação do trabalho prestado, em portfólio, site ou redes sociais, salvo manifestação expressa em contrário.")

    clause(doc, "8. SIGILO E PROPRIEDADE INTELECTUAL")
    body(doc, "8.1. Ambas as partes manterão sigilo sobre as informações trocadas em razão deste contrato.")
    body(doc, "8.2. Todo material, conteúdo e metodologia desenvolvidos permanecem sob propriedade intelectual da STUDIO, sendo concedido à CLIENTE apenas o direito de uso pessoal e intransferível. É vedada a reprodução, modificação, revenda ou comercialização sem autorização expressa e por escrito da STUDIO.")

    clause(doc, "9. DISPOSIÇÕES GERAIS")
    body(doc, "9.1. As partes são independentes, sem vínculo trabalhista, societário ou de representação.")
    body(doc, "9.2. Este contrato substitui qualquer acordo anterior entre as partes com o mesmo objeto.")
    body(doc, "9.3. As partes elegem a CAMEC BRASIL para a administração de conflitos (arbitragem, mediação e conciliação), com resolução virtual pelo site https://camecbrasil.com.br/novosite/.")
    body(doc, "9.4. As partes reconhecem como válida e plenamente eficaz a contratação por meio eletrônico e digital, ainda que com assinatura eletrônica fora dos padrões ICP-Brasil, nos termos do art. 107 do Código Civil.")

    body(doc, "São Paulo, {data}.")
    doc.add_paragraph()
    doc.add_paragraph()

    _signatures(doc)

    doc.save(out_path)


def _sig_line(doc, *linhas):
    for txt in linhas:
        p = doc.add_paragraph(txt)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _signatures(doc):
    _sig_line(doc, "_" * 40, "{nome_completo}", "CONTRATANTE")
    _sig_line(doc, "_" * 40, "Studio Sal LTDA — Giulia Salvatore Tebet Moreira", "CONTRATADA")
    # Duas testemunhas (CPC art. 784, IV — sem elas o instrumento não é título
    # executivo extrajudicial). Slots em branco, sem nomes hard-coded.
    _sig_line(doc, "_" * 40, "Nome:", "CPF:", "TESTEMUNHA 1")
    _sig_line(doc, "_" * 40, "Nome:", "CPF:", "TESTEMUNHA 2")


# ---- Serviços por tipo (das descrições reais dos contratos) ----

def servicos_mentoria(doc):
    bullet(doc, "Encontros em grupo remotos (on-line), com periodicidade semanal e datas pré-estabelecidas, avisados com antecedência mínima de 72 horas; gravações disponíveis na plataforma durante a vigência.")
    bullet(doc, "Acesso a plataforma exclusiva com aulas gravadas, mediante login e senha pessoal e intransferível, fornecidos por e-mail ou WhatsApp em até 15 dias do início do contrato.")
    bullet(doc, "Materiais auxiliares, estruturas de organização para criação de conteúdo digital, modelos e orientações, disponibilizados ao longo da vigência.")
    bullet(doc, "Suporte via WhatsApp em horário comercial, com resposta em até 48 horas úteis.")


def servicos_assessoria(doc):
    bullet(doc, "Planejamento, orientação e suporte estratégico nos canais digitais da CLIENTE.")
    bullet(doc, "Elaboração de até 2 (duas) postagens semanais, em formato de carrossel ou vídeo, para 1 (uma) rede social, conforme cronograma mensal aprovado.")
    bullet(doc, "1 (uma) reunião mensal de estratégia para alinhamento de cronograma, direcionamentos e prioridades.")
    bullet(doc, "1 (uma) reunião mensal de posicionamento, para orientação de marca pessoal e construção de narrativa.")
    bullet(doc, "Treinamento de vídeo, com orientações para gravação de conteúdo com autenticidade e naturalidade.")
    bullet(doc, "Treinamento de uso das principais plataformas (Instagram, LinkedIn, ferramentas de edição e agendamento).")
    bullet(doc, "Elaboração criativa de 1 (uma) newsletter por mês, incluindo design e conteúdo escrito.")
    bullet(doc, "Atendimento contínuo para dúvidas operacionais e estratégicas da comunicação contratada.")
    body(doc, "Não é responsabilidade da STUDIO: ativações em meios físicos/digitais não previstas, criação de websites, gestão de outras redes sociais ou os resultados comerciais obtidos. Serviços extras serão orçados à parte.")


def servicos_branding(doc):
    body(doc, "O trabalho contempla três etapas:")
    bullet(doc, "Etapa ALMA — entrevista imersiva para coleta de informações, análise do nicho (até 5 concorrentes/inspirações) e diagnóstico da comunicação atual.")
    bullet(doc, "Etapa MENTE — definição dos pilares estratégicos da marca: história, personalidade, poderes, propósito e tagline.")
    bullet(doc, "Etapa CORPO — identidade visual e verbal completa: linha editorial, diretrizes de comunicação verbal, logotipo e símbolo reduzido, paleta de cores, tipografia, moodboard e templates para Instagram (Canva).")
    body(doc, "Ao final, a STUDIO entregará o manual básico de implementação de marca com orientações para aplicação nas redes sociais.")
    body(doc, "Não é responsabilidade da STUDIO: execução de ativações/posts, campanhas de marketing, estratégias de venda, criação de websites/outras redes, ou os resultados obtidos. Serviços adicionais serão contratados à parte.")


build(
    "Contrato de Mentoria Individual",
    "Serviço de Mentoria Pessoal",
    'A STUDIO prestará à CLIENTE o serviço de mentoria denominado "{produto}", '
    "com foco em desenvolvimento pessoal e fortalecimento da marca pessoal da "
    "CLIENTE, que deverá participar ativamente do desenvolvimento do seu material.",
    servicos_mentoria,
    os.path.join(OUT_DIR, "mentoria.docx"),
    nao_participacao=True,
)

build(
    "Contrato de Prestação de Serviços de Assessoria",
    "Assessoria estratégica de posicionamento e comunicação digital",
    'A STUDIO prestará à CLIENTE os serviços contínuos de assessoria '
    'estratégica de posicionamento e comunicação digital, no formato "{produto}", '
    "conforme o escopo detalhado na Cláusula 2ª.",
    servicos_assessoria,
    os.path.join(OUT_DIR, "assessoria.docx"),
)

build(
    "Contrato de Prestação de Serviços de Branding Pessoal",
    "Desenvolvimento de branding pessoal (etapas Alma, Mente e Corpo)",
    'A STUDIO desenvolverá para a CLIENTE o projeto de branding pessoal '
    '"{produto}", compreendendo posicionamento de marca e identidade visual e '
    "verbal, conforme as três etapas detalhadas na Cláusula 2ª.",
    servicos_branding,
    os.path.join(OUT_DIR, "branding_pessoal.docx"),
)

print("done — 3 templates em", OUT_DIR)
